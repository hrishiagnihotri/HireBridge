dummy ={
    'c_name': 'string',
    'c_phone': '123456789',
    'c_email': 'hi@g.com',
    'c_skills': 'c++,python',
    'c_deg': 'B.E CSE | KLS Vishwanathrao Deshpande Institute of Technology',
    'c_degmarks': '9',
    'c_degbatch': '2025',
    'c_coll': 'abc Pu college',
    'c_collmarks': '76',
    'c_collbatch': '2020', 
    'c_school': 'abc state school', 
    'c_schoolmarks': '46', 
    'c_schoolbatch': '2016', 
    'c_int': 'haegl (2014-2018)', 'c_intdes': 'hello from haegl',
    'c_proj1': 'a','c_projdes1': 'abc',
    'c_proj2': 'b','c_projdes2': 'bcd', 
    'c_cert1': 'azure','c_certdes1': 'hi from azure', 
    'c_cert2': 'micro','c_certdes2': 'hi from micro'
       }

import os
from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



def add_bottom_border(paragraph):
    p = paragraph._p  # Access to the <w:p> element
    pPr = p.get_or_add_pPr()

    # Create <w:pBdr> if it doesn't exist
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    # Set border attributes
    bottom.set(qn('w:val'), 'single')  # Border style
    bottom.set(qn('w:sz'), '8')        # Size (1/8 pt × 8 = 1 pt)
    bottom.set(qn('w:space'), '1')     # Space between text and border
    bottom.set(qn('w:color'), '000000')  # Black color

    # Add bottom to pBdr and pBdr to pPr
    pBdr.append(bottom)
    pPr.append(pBdr)

def addheadline(doc,stri):
    headline = doc.add_paragraph(stri)
    run = headline.runs[0]
    run.font.name = 'Aptos SemiBold'
    add_bottom_border(headline)

def makeresumeoutofit(data,foldername):
    doc = Document()
    skills = data['c_skills'].split(',')
    style = doc.styles['Normal']
    style.font.name = 'Aptos'
    style.font.size = Pt(12)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    title = doc.add_heading(data['c_name'].capitalize(),1)
    run=title.runs[0]
    run.font.size = Pt(22)
    title.paragraph_format.space_before = Pt(0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contacts = doc.add_paragraph(f'Email: {data["c_email"]} | Phone: {data["c_phone"]}')
    contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contacts.runs[0]
    run.font.name = 'Aptos SemiBold'
    # run.bold = True

    addheadline(doc,'SKILLS')
    for s in skills:
        point = doc.add_paragraph(s.capitalize(),style='ListBullet')
        point.runs[0].font.name = 'Aptos SemiBold'
    addheadline(doc,'EDUCATION')

    count=0
    if data['c_coll'] != '':
        count+=1
    if data['c_school'] != '':
        count+=1

    edu_table = doc.add_table(1+count,3)


    edu_table.cell(0,0)._element.clear_content()
    a=edu_table.cell(0,0).add_paragraph(data['c_deg'],style='ListBullet')
    edu_table.cell(0,2).width = Cm(1.25)
    edu_table.cell(0,1).width = Cm(5.3)
    edu_table.cell(0,0).width = Cm(51)
    edu_table.cell(0,1).text = "CGPA: "+data['c_degmarks']
    edu_table.cell(0,2).text = f'({data['c_degbatch']})'
    i=1
    if data['c_coll'] != '':
        edu_table.cell(i,0)._element.clear_content()
        edu_table.cell(i,0).add_paragraph(data['c_coll'],style='ListBullet')
        edu_table.cell(i,1).text = data['c_collmarks']+"%"
        edu_table.cell(i,2).text = f'({data['c_collbatch']})'
        i+=1

    if data['c_school'] != '':
        edu_table.cell(i,0)._element.clear_content()
        edu_table.cell(i,0).add_paragraph(data['c_school'],style='ListBullet')
        edu_table.cell(i,1).text = data['c_schoolmarks']+'%'
        edu_table.cell(i,2).text = f'({data['c_schoolbatch']})'
        i+=1

    if data['c_int']!= '':
        addheadline(doc,"INTERNSHIP")
        int_title=doc.add_paragraph(data['c_int'].upper())
        int_title.runs[0].bold =True
        int_title.paragraph_format.space_after = Pt(2)
        int_desc = doc.add_paragraph(data['c_intdes'].capitalize())


    if data['c_proj1']!= '' or data['c_proj2']!= '':
        addheadline(doc,"PROJECTS")
        if data['c_proj1']!= '':
            proj1_title=doc.add_paragraph(data['c_proj1'].upper())
            proj1_title.runs[0].bold =True
            proj1_title.paragraph_format.space_after = Pt(2)
            proj1_desc = doc.add_paragraph(data['c_projdes1'].capitalize())
        if data['c_proj2']!= '':
            proj2_title=doc.add_paragraph(data['c_proj2'].upper())
            proj2_title.runs[0].bold =True
            proj2_title.paragraph_format.space_after = Pt(2)
            proj2_desc = doc.add_paragraph(data['c_projdes2'].capitalize())

    if data['c_cert1']!= '' or data['c_cert2']!= '':
        addheadline(doc,"CERTIFICATIONS")
        if data['c_cert1']!= '':
            cert1_title=doc.add_paragraph(data['c_cert1'].upper())
            cert1_title.runs[0].bold =True
            cert1_title.paragraph_format.space_after = Pt(2)
            cert1_desc = doc.add_paragraph(data['c_certdes1'].capitalize())
        if data['c_cert2']!= '':
            cert2_title=doc.add_paragraph(data['c_cert2'].upper())
            cert2_title.runs[0].bold =True
            cert2_title.paragraph_format.space_after = Pt(2)
            cert2_desc = doc.add_paragraph(data['c_certdes2'].capitalize())
    foldername =f"userdata\\{foldername}"
    if not os.path.exists(foldername):
        os.makedirs(foldername)
    save_path = os.path.join(foldername, "resume.docx")
    doc.save(save_path)
    print('generated')
    return save_path

